"""文档解析 — 使用 PyPDF2 + python-docx 逐页/逐段提取文本，支持部分容错

对齐 ARCHITECTURE.md §4.7:
- 单页/单段失败跳过并记录 warning
- < 20% 失败 → 继续（记录 warning）
- 20%~50% 失败 → partial_failed
- > 50% 失败 → failed

对齐 ROADMAP.md §8.7（Chunk 元数据增强）：
- DOCX 标题样式自动转换为 Markdown # 标记，使 chunker 的标题检测跨格式统一
"""

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from PyPDF2 import PdfReader

logger = logging.getLogger(__name__)

# Word 内置标题样式名映射 -> Markdown 标题层级
_WORD_HEADING_PATTERN = re.compile(r'^Heading\s*(\d+)', re.IGNORECASE)
_WORD_TITLE_PATTERNS = re.compile(r'^(Title|Subtitle)$', re.IGNORECASE)


@dataclass
class ParsedPage:
    """单页解析结果"""
    page_number: int
    content: str
    success: bool = True
    error: str | None = None


@dataclass
class ParseResult:
    """文档解析聚合结果"""
    pages: list[ParsedPage] = field(default_factory=list)
    total_pages: int = 0
    failed_pages: int = 0
    source_type: str = ""  # pdf / docx / md / txt，用于日志单位判断

    @property
    def failure_rate(self) -> float:
        if self.total_pages == 0:
            return 1.0  # 空文档视为全部失败
        return self.failed_pages / self.total_pages

    @property
    def full_text(self) -> str:
        """拼接所有成功页面的文本"""
        return "\n\n".join(p.content for p in self.pages if p.success)

    @property
    def warnings(self) -> list[str]:
        """收集所有失败页面的警告信息"""
        return [
            f"第{p.page_number}页: {p.error}"
            for p in self.pages
            if not p.success and p.error
        ]


def parse_document(file_path: str, file_type: str | None = None) -> ParseResult:
    """解析文档主入口，根据文件类型分发到对应解析器。

    Args:
        file_path: 文档文件绝对路径
        file_type: 文件类型（pdf/docx/md/txt），为 None 时从扩展名推断

    Returns:
        ParseResult: 包含逐页/逐段解析结果和容错统计
    """
    path = Path(file_path)
    if not path.exists():
        return ParseResult(
            pages=[ParsedPage(page_number=1, content="", success=False, error=f"文件不存在: {file_path}")],
            total_pages=1,
            failed_pages=1,
            source_type=file_type or "",
        )

    if file_type is None:
        file_type = path.suffix.lower().lstrip(".")

    try:
        if file_type == "pdf":
            result = _parse_pdf(file_path)
        elif file_type == "docx":
            result = _parse_docx(file_path)
        elif file_type in ("md", "txt"):
            result = _parse_text(file_path)
        else:
            return ParseResult(
                pages=[ParsedPage(page_number=1, content="", success=False, error=f"不支持的文件类型: {file_type}")],
                total_pages=1,
                failed_pages=1,
                source_type=file_type,
            )
        result.source_type = file_type
        return result
    except Exception as e:
        logger.exception(f"文档解析异常: {file_path}")
        return ParseResult(
            pages=[ParsedPage(page_number=1, content="", success=False, error=str(e))],
            total_pages=1,
            failed_pages=1,
            source_type=file_type or "",
        )


def _parse_pdf(file_path: str) -> ParseResult:
    """使用 PyPDF2 逐页解析 PDF，单页失败跳过并记录"""
    try:
        reader = PdfReader(file_path)
    except Exception as e:
        return ParseResult(
            pages=[ParsedPage(page_number=1, content="", success=False, error=str(e))],
            total_pages=1,
            failed_pages=1,
        )

    pages: list[ParsedPage] = []
    failed = 0

    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text()
            if text and text.strip():
                pages.append(ParsedPage(page_number=i + 1, content=text.strip()))
            else:
                pages.append(ParsedPage(
                    page_number=i + 1, content="",
                    success=False, error="页面无文本或文本为空"
                ))
                failed += 1
        except Exception as e:
            pages.append(ParsedPage(
                page_number=i + 1, content="",
                success=False, error=f"页面解析异常: {e}"
            ))
            failed += 1

    total = len(reader.pages)
    return ParseResult(pages=pages, total_pages=total, failed_pages=failed)


def _docx_heading_to_markdown(paragraph) -> str | None:
    """检测 Word 段落是否为标题样式，返回对应的 Markdown # 前缀文本。

    对齐 ROADMAP.md §8.7：将 DOCX 标题样式转换为 Markdown 标记，
    使 chunker.py 的 detect_sections() 可跨 MD/DOCX 统一检测。

    防御性设计：MagicMock 等非真实对象会导致属性访问异常，
    此时返回 None 降级为普通文本提取。

    Args:
        paragraph: python-docx Paragraph 对象

    Returns:
        带 # 前缀的标题文本，或 None（非标题段落/异常）
    """
    try:
        style = paragraph.style
        if style is None:
            return None

        text = paragraph.text
        if not text or not text.strip():
            return None

        # 检查段落样式（优先）
        style_name = style.name or ""
        m = _WORD_HEADING_PATTERN.match(style_name)
        if m:
            level = int(m.group(1))
            if 1 <= level <= 6:
                return f"{'#' * level} {text.strip()}"

        # Title/Subtitle → # / ##
        if _WORD_TITLE_PATTERNS.match(style_name):
            if style_name.lower() == "title":
                return f"# {text.strip()}"
            else:
                return f"## {text.strip()}"

        # 检查大纲级别（Word 内置段落属性 outline_level，如 outlineLvl）
        try:
            outline_lvl = paragraph.paragraph_format.outline_level
            if outline_lvl is not None and 0 <= outline_lvl <= 5:
                return f"{'#' * (outline_lvl + 1)} {text.strip()}"
        except (AttributeError, ValueError, TypeError):
            pass

        # 检查段落样式类型是否为 HEADING
        if style.type == WD_STYLE_TYPE.PARAGRAPH and hasattr(style, 'base_style'):
            try:
                base = style.base_style
                if base is not None:
                    base_name = base.name or ""
                    m2 = _WORD_HEADING_PATTERN.match(base_name)
                    if m2:
                        level = int(m2.group(1))
                        if 1 <= level <= 6:
                            return f"{'#' * level} {text.strip()}"
            except (AttributeError, ValueError, TypeError):
                pass

        return None
    except (AttributeError, TypeError, ValueError):
        # MagicMock 等非真实对象 → 降至普通文本
        return None


def _parse_docx(file_path: str) -> ParseResult:
    """使用 python-docx 解析 DOCX，逐段提取并容错（对齐 PDF 逐页容错粒度）。

    DOCX 标题样式自动转换为 Markdown # 标记，使 chunker.py 的章节检测
    跨 MD/DOCX 格式统一工作（对齐 ROADMAP.md §8.7）。
    """
    try:
        doc = DocxDocument(file_path)
    except Exception as e:
        return ParseResult(
            pages=[ParsedPage(page_number=1, content="", success=False, error=str(e))],
            total_pages=1,
            failed_pages=1,
        )

    if not doc.paragraphs:
        return ParseResult(
            pages=[ParsedPage(page_number=1, content="", success=False, error="文档无段落内容")],
            total_pages=1,
            failed_pages=1,
        )

    pages: list[ParsedPage] = []
    failed = 0

    for i, p in enumerate(doc.paragraphs):
        try:
            # 检测标题样式，转换为 Markdown 标记（§8.7）
            heading_text = _docx_heading_to_markdown(p)
            if heading_text is not None:
                pages.append(ParsedPage(page_number=i + 1, content=heading_text))
            else:
                text = p.text
                if text and text.strip():
                    pages.append(ParsedPage(page_number=i + 1, content=text.strip()))
        except Exception as e:
            logger.warning(f"DOCX 第{i+1}段解析失败: {e}")
            pages.append(ParsedPage(
                page_number=i + 1, content="",
                success=False, error=f"段落解析异常: {e}"
            ))
            failed += 1

    total = len(doc.paragraphs)

    if not pages:
        return ParseResult(
            pages=[ParsedPage(page_number=1, content="", success=False, error="文档无有效文本内容")],
            total_pages=total,
            failed_pages=total,
        )

    return ParseResult(pages=pages, total_pages=total, failed_pages=failed)


def _parse_text(file_path: str) -> ParseResult:
    """解析纯文本文件（md/txt），统一 UTF-8 读取"""
    try:
        content = Path(file_path).read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            content = Path(file_path).read_text(encoding="gbk")
        except Exception as e:
            return ParseResult(
                pages=[ParsedPage(page_number=1, content="", success=False, error=f"编码错误: {e}")],
                total_pages=1,
                failed_pages=1,
            )
    except Exception as e:
        return ParseResult(
            pages=[ParsedPage(page_number=1, content="", success=False, error=str(e))],
            total_pages=1,
            failed_pages=1,
        )

    if not content.strip():
        return ParseResult(
            pages=[ParsedPage(page_number=1, content="", success=False, error="文件内容为空")],
            total_pages=1,
            failed_pages=1,
        )

    page = ParsedPage(page_number=1, content=content.strip())
    return ParseResult(pages=[page], total_pages=1, failed_pages=0)
